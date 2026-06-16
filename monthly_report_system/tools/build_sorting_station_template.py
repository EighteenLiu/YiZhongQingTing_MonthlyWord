"""Build a cleaned placeholder template for garbage sorting station reports."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = PROJECT_DIR / "templates"
SRC_CANDIDATES = [
    TEMPLATE_DIR / "sorting_station_template_repaired.docx",
    TEMPLATE_DIR / "西城区20XX年X月生活垃圾分类驿站检查通报_模板.docx",
]
SRC = next((path for path in SRC_CANDIDATES if path.exists()), SRC_CANDIDATES[0])
OUT = TEMPLATE_DIR / "garbage_sorting_station_report_template.docx"
CN_NUMS = "一二三四五六七八九十"


def strip_generated_prefix(text: str) -> str:
    return re.sub(r"^(?:[（(]\d+[）)]\s*)+", "", text.strip())


def clean_key(text: str) -> str:
    text = strip_generated_prefix(text)
    text = re.sub(r"\{\s*problem\s*\}", "", text.strip(), flags=re.I)
    text = re.sub(r"^\s*\d+[.．、]\s*", "", text)
    text = re.split(r"[:：﹕꞉]", text, maxsplit=1)[0]
    text = re.sub(r'[{}\s\\/:*?"<>|（）()，,。；;、.．]+', "_", text)
    return text.strip("_") or "未命名点位"


def normalize_street_heading(text: str) -> str:
    stripped = strip_generated_prefix(text)
    if len(stripped) >= 3 and stripped[0] == "（" and "）" in stripped[:5]:
        close = stripped.find("）")
        if all(ch in CN_NUMS for ch in stripped[1:close]):
            return stripped[close + 1 :].strip()
    return stripped


def replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def remove_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    num_pr = p_pr.numPr
    if num_pr is not None:
        p_pr.remove(num_pr)


def problem_key(text: str) -> str:
    start = text.find("{{")
    end = text.find("_问题列表}}", start)
    if start == -1 or end == -1:
        return ""
    return text[start + 2 : end]


def image_key(text: str) -> str:
    match = re.search(r"\{\{(.+)_图片\d+\}\}", text)
    return match.group(1) if match else ""


IMAGE_PLACEHOLDER_ROWS = 3
IMAGE_PLACEHOLDER_COLS = 2
IMAGE_PLACEHOLDER_WIDTH_CM = 7.87
IMAGE_PLACEHOLDER_HEIGHT_CM = 4.43


def insert_table_after(document, paragraph, rows: int = IMAGE_PLACEHOLDER_ROWS, cols: int = IMAGE_PLACEHOLDER_COLS):
    table = document.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)  # noqa: SLF001
    return table


def insert_table_after_table(document, previous_table, rows: int = IMAGE_PLACEHOLDER_ROWS, cols: int = IMAGE_PLACEHOLDER_COLS):
    table = document.add_table(rows=rows, cols=cols)
    previous_table._tbl.addnext(table._tbl)  # noqa: SLF001
    return table


def set_cell(cell, text: str) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "楷体_GB2312"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体_GB2312")  # noqa: SLF001
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(102, 102, 102)


def set_image_placeholder_table(table, key: str) -> None:
    set_table_borders(table)
    set_widths(table)
    idx = 1
    for row in table.rows:
        for cell in row.cells:
            set_cell(cell, "{{" + f"{key}_图片{idx}" + "}}")
            idx += 1


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr  # noqa: SLF001
    old = tbl_pr.first_child_found_in("w:tblBorders")
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "dashed")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7B7B7")
        borders.append(element)


def set_widths(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = Cm(IMAGE_PLACEHOLDER_WIDTH_CM)
    height = Cm(IMAGE_PLACEHOLDER_HEIGHT_CM)
    dxa = str(int(width.emu / 635))
    table_width_dxa = str(int(width.emu / 635) * IMAGE_PLACEHOLDER_COLS)
    tbl_pr = table._tbl.tblPr  # noqa: SLF001
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), table_width_dxa)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid  # noqa: SLF001
    if tbl_grid is not None:
        table._tbl.remove(tbl_grid)  # noqa: SLF001
    tbl_grid = OxmlElement("w:tblGrid")
    for _ in range(IMAGE_PLACEHOLDER_COLS):
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), dxa)
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)  # noqa: SLF001
    for row in table.rows:
        row.height = height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), dxa)
            tc_w.set(qn("w:type"), "dxa")


def relayout_existing_image_tables(document) -> int:
    changed = 0
    for table in list(document.tables):
        if not table.rows:
            continue
        key = image_key(table.rows[0].cells[0].text.strip())
        if not key:
            continue
        new_table = insert_table_after_table(document, table)
        set_image_placeholder_table(new_table, key)
        table._tbl.getparent().remove(table._tbl)  # noqa: SLF001
        changed += 1
    return changed


def main() -> None:
    document = Document(str(SRC))
    current_street = ""
    street_counter = 0
    seen_streets: list[str] = []
    problem_paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        clean_text = strip_generated_prefix(text)
        if "西城区生活垃圾分类驿站检查通报" in clean_text and "{" not in clean_text:
            remove_paragraph_numbering(paragraph)
            replace_paragraph_text(paragraph, "西城区{{报告年份}}年{{报告月份}}生活垃圾分类驿站检查通报")
            continue
        if clean_text.startswith("2026年{}至2026年{}") or clean_text.startswith("{{检查开始日期}}至{{检查结束日期}}"):
            remove_paragraph_numbering(paragraph)
            replace_paragraph_text(
                paragraph,
                "{{检查开始日期}}至{{检查结束日期}}对本区{{街道数量}}个街道生活垃圾分类驿站进行了{{检查站次}}个次检查。检查发现需通报问题站次{{问题站次}}个。",
            )
            continue
        if clean_text.startswith("附件："):
            remove_paragraph_numbering(paragraph)
            replace_paragraph_text(paragraph, "附件：{{检查开始日期}}至{{检查结束日期}}生活垃圾分类驿站各项指标问题数")
            continue

        if "{problem}" in clean_text:
            remove_paragraph_numbering(paragraph)
            point_key = clean_key(clean_text)
            key_prefix = f"{current_street}_{point_key}".strip("_")
            placeholder = "{{" + f"{key_prefix}_问题列表" + "}}"
            replace_paragraph_text(paragraph, clean_text.replace("{problem}", placeholder))
            problem_paragraphs.append(paragraph)
            continue

        is_section = clean_text in {"一、总体情况", "二、各街道情况"}
        is_point = "{problem}" in clean_text or "：" in clean_text or ":" in clean_text or re.match(r"^\d+[.．、]", clean_text)
        if not is_section and not is_point:
            remove_paragraph_numbering(paragraph)
            street = normalize_street_heading(clean_text)
            current_street = street
            if street and street not in seen_streets:
                street_counter += 1
                seen_streets.append(street)
            replace_paragraph_text(paragraph, street)

    for paragraph in problem_paragraphs:
        key = problem_key(paragraph.text)
        if not key:
            continue
        table = insert_table_after(document, paragraph)
        set_image_placeholder_table(table, key)

    relayout_count = relayout_existing_image_tables(document)

    for table in document.tables:
        if table.rows:
            for cell in table.rows[0].cells:
                cell.text = cell.text.strip()
        if len(table.rows) > 1 and table.rows[0].cells[0].text.strip() == "街道名称":
            table.rows[1].cells[0].text = "总计"

    document.save(OUT)
    print(f"saved {OUT}")
    print(f"problem paragraphs {len(problem_paragraphs)}")
    print(f"relayout image tables {relayout_count}")


if __name__ == "__main__":
    main()
