from __future__ import annotations

from collections import OrderedDict
import sys
from pathlib import Path

import pandas as pd
import xlrd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
INPUT_DIR = ROOT / "input"
MODULE_DIR = INPUT_DIR / "module"
OUTPUT_DIR = ROOT / "output"

TEMPLATE_PATH = MODULE_DIR / "环卫停车场检查报告_Jinja模板.docx"
REPORT_SOURCE_PATH = INPUT_DIR / "202606YZQT.xls"
REQ_DOC_PATH = OUTPUT_DIR / "环卫停车场模板修改需求文档.docx"
SAMPLE_REPORT_PATH = OUTPUT_DIR / "西城区2026年6月环卫停车场检查报告.docx"

PARKING_ITEMS = ["作业安全", "设施设备", "消防安全", "场内环境", "场内容器"]


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # noqa: SLF001
    new_para = paragraph._parent.add_paragraph()  # noqa: SLF001
    new_para._p.getparent().remove(new_para._p)  # noqa: SLF001
    new_p.addnext(new_para._p)  # noqa: SLF001
    if text:
        new_para.text = text
    return new_para


def set_run_font(run, size: int, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")  # noqa: SLF001
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")  # noqa: SLF001
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")  # noqa: SLF001
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def rewrite_template() -> None:
    doc = Document(str(TEMPLATE_PATH))
    while len(doc.paragraphs) < 20:
        doc.add_paragraph()

    texts = [
        '{{ title | default("西城区环卫作业停车场检查报告") }}',
        '{{ report_range_text }}西城区城管委固废科对{{ site_names_text }}环卫作业停车场进行现场检查。',
        '{{ department | default("北京市西城区城市管理委员会") }}',
        "{%p for site in sites %}",
        "{{ site.index_cn }}、{{ site.name }}",
        "{%p if site.photo_rows %}",
        "{%p for row in site.photo_rows %}",
        "{{ row.left }}{{ row.right }}",
        "{%p endfor %}",
        "{%p endif %}",
        "{%p for item in site.items %}",
        '{{ loop.index }}.{{ item.name }}：{{ item.result_text | default("无问题。") }}',
        "{%p if item.photo_rows %}",
        "{%p for row in item.photo_rows %}",
        "{{ row.left }}{{ row.right }}",
        "{%p endfor %}",
        "{%p endif %}",
        "{%p endfor %}",
        '检查问题：{{ site.issue_summary | default("无问题。") }}',
        "{%p endfor %}",
        '{{ closing_text | default("所查点位发现问题，所属单位立即整改，排除安全隐患/加强日常管理/做好培训工作。所查点位未发现问题，望各单位继续保持，共建美丽西城。") }}',
        '{{ department | default("北京市西城区城市管理委员会") }}',
        "{{ issue_date_text }}",
    ]

    # Preserve the existing first three paragraphs and replace them in place.
    for idx, text in enumerate(texts):
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
        else:
            doc.add_paragraph(text)

    doc.save(str(TEMPLATE_PATH))


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=3, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 16, bold=True, color="000000")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, before=12, after=6, line=1.1)
        run = p.add_run(text)
        set_run_font(run, 13, bold=True, color="2E74B5")
    else:
        style_paragraph(p, before=8, after=4, line=1.1)
        run = p.add_run(text)
        set_run_font(run, 12, bold=True, color="1F4D78")


def add_body(doc: Document, text: str, after: int = 6) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=after, line=1.15)
    run = p.add_run(text)
    set_run_font(run, 11, color="000000")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    for idx, (header, width) in enumerate(zip(headers, widths_cm)):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.width = Cm(width)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, before=0, after=0, line=1.0)
            for run in p.runs:
                set_run_font(run, 10, bold=True, color="FFFFFF")
        tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "2E74B5")
        tc_pr.append(shd)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            cell.width = Cm(widths_cm[idx])
            for p in cell.paragraphs:
                style_paragraph(p, before=0, after=0, line=1.1)
                for run in p.runs:
                    set_run_font(run, 10, color="000000")


def build_requirement_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    add_title(doc, "环卫停车场模板修改需求文档")
    add_body(doc, "目的：根据2026年5月月报与2026年6月台账，统一修订环卫停车场 Jinja 模板的标题层级、指标顺序和问题展示规则。")

    add_heading(doc, "输入文件", 1)
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["环卫停车场检查报告_Jinja模板.docx", "当前模板基线，需改成适配6月台账的新版模板"],
            ["2026年4月20日-2026年5月19日环卫停车场检查记录.docx", "5月成稿，作为标题顺序与行文样式参照"],
            ["202606YZQT.xls", "6月台账，作为实际数据源"],
        ],
        [4.5, 10.0],
    )

    add_heading(doc, "生成规则", 1)
    rules = [
        "一级标题按4级点位输出，即每个停车场点位单独成节，编号按“一、二、三……”连续排列。",
        "二级标题按5月成稿中的指标顺序输出：作业安全、设施设备、消防安全、场内环境、场内容器。",
        "只输出台账中实际检查过的指标；没有对应记录的指标整项不写。",
        "二级标题冒号后的内容直接参考“具体问题”列：该列为“无问题”或其他无问题表述时，统一写“无问题”；否则写具体问题原文。",
        "指标如带图片，图片紧随对应指标后展示；点位级图片可在点位标题后统一展示。",
        "前缀编号“1.”“2.”等必须按每个点位内部连续，不跳号。",
    ]
    for rule in rules:
        add_body(doc, f"• {rule}")

    add_heading(doc, "输出结构", 1)
    add_table(
        doc,
        ["层级", "来源字段", "展示规则"],
        [
            ["一级标题", "4级点位", "保留点位名，作为单个停车场小节标题"],
            ["二级标题", "2级指标/业务指标", "按5月月报固定顺序排序，仅显示实际检查项"],
            ["冒号后正文", "具体问题", "无问题类表述统一写“无问题”，否则原文输出"],
            ["图片", "图片列", "指标图片在指标后展示，点位图片在点位标题后展示"],
        ],
        [2.5, 4.5, 8.0],
    )

    add_heading(doc, "验收口径", 1)
    checks = [
        "打开模板后仍是可渲染的 Jinja 文档，循环与条件块完整。",
        "6月停车场样例生成后，3个点位的指标顺序与5月成稿一致。",
        "无检查项的指标不应出现在正文中。",
        "“无问题”“无”“正常”等无问题表述应统一落成“无问题”。",
        "存在问题时，正文应保留具体问题原文，例如“问题”“无容器”。",
    ]
    for check in checks:
        add_body(doc, f"□ {check}")

    doc.save(str(REQ_DOC_PATH))


def build_sample_report() -> None:
    from core.data_cleaner import attach_images_to_records, clean_dataframe
    from core.excel_reader import read_excel_table
    from core.image_extractor import extract_images_by_row
    from core.parking_station_report_generator import generate_parking_station_report

    temp_dir = OUTPUT_DIR / "_sample_build"
    temp_dir.mkdir(parents=True, exist_ok=True)
    df, actual_excel_path = read_excel_table(REPORT_SOURCE_PATH, temp_dir)
    clean = clean_dataframe(df, "parking_station")
    images_by_row, _ = extract_images_by_row(actual_excel_path, temp_dir / "images", max_images_per_row=15)
    records, _ = attach_images_to_records(clean.records, images_by_row, max_images_per_record=15, warn_on_missing=False, warn_on_partial=False, ignored_row_numbers=clean.filtered_row_numbers)
    generate_parking_station_report(
        records=records,
        output_path=SAMPLE_REPORT_PATH,
        title="西城区环卫作业停车场检查报告",
        start_date="2026年5月20日",
        end_date="2026年6月19日",
        report_month="2026年6月",
        template_path=TEMPLATE_PATH,
    )


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    rewrite_template()
    build_requirement_doc()
    build_sample_report()


if __name__ == "__main__":
    main()
