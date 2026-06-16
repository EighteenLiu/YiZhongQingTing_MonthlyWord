"""Word 月报生成，支持基于上传模板写入内容。"""

from __future__ import annotations

import re
from collections import OrderedDict
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config.indicators import REPORT_STYLE, ReportStyle, indicators_for


CHINESE_NUMBERS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
A4_WIDTH_CM = 21
A4_HEIGHT_CM = 29.7
PLAIN_PROBLEM_TEXTS = {
    "无问题",
    "未营业",
    "人员不在，未营业",
    "未发现点位",
    "未见点位",
    "点位已撤",
    "未见点位或点位已撤",
    "无点位或点位已撤",
}


@dataclass
class TextFormat:
    """从模板样本文字中提取出的段落和字符格式。"""

    style_name: str | None = None
    p_pr: Any | None = None
    r_pr: Any | None = None


@dataclass
class TemplateFormat:
    """交投点模板中可复用的格式样本。"""

    title: TextFormat | None = None
    heading1: TextFormat | None = None
    heading2: TextFormat | None = None
    heading3: TextFormat | None = None
    body: TextFormat | None = None
    table_header: TextFormat | None = None
    table_body: TextFormat | None = None
    table_style_name: str | None = None


@dataclass
class TemplateSections:
    """模板首页正文节与附件节的页面设置。"""

    first_section_pr: Any | None = None
    last_section_pr: Any | None = None


def _capture_template_sections(document: Document) -> TemplateSections:
    """保存模板首节和末节页面设置，避免清空正文后只剩最后一节。"""

    section_properties = [deepcopy(section._sectPr) for section in document.sections]  # noqa: SLF001
    if not section_properties:
        return TemplateSections()
    return TemplateSections(first_section_pr=section_properties[0], last_section_pr=section_properties[-1])


def _replace_body_section_properties(document: Document, section_pr: Any | None) -> None:
    """替换正文末尾 sectPr。"""

    body = document._body._element  # noqa: SLF001
    existing = body.sectPr
    if existing is not None:
        body.remove(existing)
    if section_pr is not None:
        body.append(deepcopy(section_pr))


def _clear_body_keep_section(document: Document, section_pr: Any | None = None) -> None:
    """清空模板正文内容，同时保留指定 sectPr、样式、页眉页脚和页面设置。"""

    body = document._body._element  # noqa: SLF001
    current_section_pr = body.sectPr
    source_section_pr = section_pr if section_pr is not None else current_section_pr
    saved_section_properties = deepcopy(source_section_pr) if source_section_pr is not None else None
    for child in list(body):
        body.remove(child)
    if saved_section_properties is not None:
        body.append(saved_section_properties)


def _capture_text_format(paragraph) -> TextFormat:
    """复制一个段落的样式、段落属性和首个有文本 run 的字符属性。"""

    run = next((item for item in paragraph.runs if item.text.strip()), None)
    if run is None and paragraph.runs:
        run = paragraph.runs[0]
    return TextFormat(
        style_name=paragraph.style.name if paragraph.style is not None else None,
        p_pr=deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None,  # noqa: SLF001
        r_pr=deepcopy(run._r.rPr) if run is not None and run._r.rPr is not None else None,  # noqa: SLF001
    )


def _is_heading1_text(text: str) -> bool:
    return bool(re.match(r"^[一二三四五六七八九十]+、", text)) or text.startswith("附件：")


def _is_heading2_text(text: str) -> bool:
    return bool(re.match(r"^（[一二三四五六七八九十]+）", text))


def _is_heading3_text(text: str) -> bool:
    return bool(re.match(r"^\d+[.．、]", text)) and "交投点" in text


def _capture_template_format(document: Document) -> TemplateFormat:
    """从上传模板正文中抓取各类文本和表格的格式样本。"""

    non_empty_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    template_format = TemplateFormat()
    if non_empty_paragraphs:
        template_format.title = _capture_text_format(non_empty_paragraphs[0])

    for paragraph in non_empty_paragraphs[1:]:
        text = paragraph.text.strip()
        if template_format.heading1 is None and _is_heading1_text(text):
            template_format.heading1 = _capture_text_format(paragraph)
            continue
        if template_format.heading2 is None and _is_heading2_text(text):
            template_format.heading2 = _capture_text_format(paragraph)
            continue
        if template_format.heading3 is None and _is_heading3_text(text):
            template_format.heading3 = _capture_text_format(paragraph)
            continue
        if template_format.body is None and not (_is_heading1_text(text) or _is_heading2_text(text) or _is_heading3_text(text)):
            template_format.body = _capture_text_format(paragraph)

    if template_format.body is None and non_empty_paragraphs:
        template_format.body = _capture_text_format(non_empty_paragraphs[-1])
    if template_format.heading1 is None:
        template_format.heading1 = template_format.body
    if template_format.heading2 is None:
        template_format.heading2 = template_format.heading1
    if template_format.heading3 is None:
        template_format.heading3 = template_format.body

    if document.tables:
        table = document.tables[0]
        template_format.table_style_name = table.style.name if table.style is not None else None
        if table.rows:
            template_format.table_header = _capture_text_format(table.rows[0].cells[0].paragraphs[0])
        if len(table.rows) > 1:
            template_format.table_body = _capture_text_format(table.rows[1].cells[0].paragraphs[0])
        else:
            template_format.table_body = template_format.table_header or template_format.body

    return template_format


def _apply_text_format(paragraph, text_format: TextFormat | None) -> None:
    """把模板段落样式应用到新段落。"""

    if text_format is None:
        return
    if text_format.style_name:
        try:
            paragraph.style = text_format.style_name
        except KeyError:
            pass
    if text_format.p_pr is not None:
        existing = paragraph._p.pPr  # noqa: SLF001
        if existing is not None:
            paragraph._p.remove(existing)  # noqa: SLF001
        paragraph._p.insert(0, deepcopy(text_format.p_pr))  # noqa: SLF001


def _apply_run_format(run, text_format: TextFormat | None) -> None:
    """把模板字符样式应用到新 run。"""

    if text_format is None or text_format.r_pr is None:
        return
    existing = run._r.rPr  # noqa: SLF001
    if existing is not None:
        run._r.remove(existing)  # noqa: SLF001
    run._r.insert(0, deepcopy(text_format.r_pr))  # noqa: SLF001


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # noqa: SLF001
    run.font.size = Pt(size_pt)
    run.bold = bold


def _set_paragraph_font(paragraph, font_name: str, size_pt: int, bold: bool = False) -> None:
    for run in paragraph.runs:
        _set_run_font(run, font_name, size_pt, bold)


def _set_cell_text(
    cell,
    text: str,
    font_name: str,
    size_pt: int,
    bold: bool = False,
    text_format: TextFormat | None = None,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_text_format(paragraph, text_format)
    run = paragraph.add_run(str(text))
    if text_format is not None:
        _apply_run_format(run, text_format)
    else:
        _set_run_font(run, font_name, size_pt, bold)


def _set_table_borders(table) -> None:
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)


def _set_a4_section(section, orientation: WD_ORIENT, style: ReportStyle) -> None:
    """设置 A4 页面方向和边距。"""

    section.orientation = orientation
    if orientation == WD_ORIENT.LANDSCAPE:
        section.page_width = Cm(A4_HEIGHT_CM)
        section.page_height = Cm(A4_WIDTH_CM)
    else:
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)
    section.top_margin = Cm(style.page_margin_cm)
    section.bottom_margin = Cm(style.page_margin_cm)
    section.left_margin = Cm(style.page_margin_cm)
    section.right_margin = Cm(style.page_margin_cm)


def _add_landscape_section(
    document: Document,
    style: ReportStyle,
    keep_template_margins: bool = False,
    template_sections: TemplateSections | None = None,
) -> None:
    """为附件统计表创建独立横向 A4 分节。"""

    previous_section = document.sections[-1]
    section = document.add_section(WD_SECTION.NEW_PAGE)
    if template_sections is not None and template_sections.last_section_pr is not None:
        _replace_body_section_properties(document, template_sections.last_section_pr)
    elif keep_template_margins:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = previous_section.page_height
        section.page_height = previous_section.page_width
        section.top_margin = previous_section.top_margin
        section.bottom_margin = previous_section.bottom_margin
        section.left_margin = previous_section.left_margin
        section.right_margin = previous_section.right_margin
    else:
        _set_a4_section(section, WD_ORIENT.LANDSCAPE, style)


def _add_title(document: Document, title: str, style: ReportStyle, template_format: TemplateFormat | None = None) -> None:
    paragraph = document.add_paragraph()
    text_format = template_format.title if template_format else None
    if text_format is not None:
        _apply_text_format(paragraph, text_format)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    if text_format is not None:
        _apply_run_format(run, text_format)
    else:
        _set_run_font(run, style.title_font, style.title_size_pt, True)


def _add_heading(document: Document, text: str, level: int, template_format: TemplateFormat | None = None) -> None:
    """按 Word 标题级别添加标题，优先使用模板中的标题样式。"""

    text_format = None
    if template_format is not None:
        text_format = {
            1: template_format.heading1,
            2: template_format.heading2,
            3: template_format.heading3,
        }.get(level)
    if text_format is not None:
        paragraph = document.add_paragraph()
        _apply_text_format(paragraph, text_format)
        run = paragraph.add_run(text)
        _apply_run_format(run, text_format)
    else:
        paragraph = document.add_heading(text, level=level)
        paragraph.paragraph_format.first_line_indent = Pt(0)


def _add_body_paragraph(
    document: Document,
    text: str,
    style: ReportStyle,
    indent: bool = True,
    template_format: TemplateFormat | None = None,
) -> None:
    paragraph = document.add_paragraph()
    text_format = template_format.body if template_format else None
    if text_format is not None:
        _apply_text_format(paragraph, text_format)
    elif indent:
        paragraph.paragraph_format.first_line_indent = Pt(style.body_size_pt * style.first_line_indent_chars)
    run = paragraph.add_run(text)
    if text_format is not None:
        _apply_run_format(run, text_format)
    else:
        _set_run_font(run, style.body_font, style.body_size_pt)


def _add_point_problem_paragraph(
    document: Document,
    heading_text: str,
    problem_text: str,
    style: ReportStyle,
    template_format: TemplateFormat | None = None,
) -> None:
    text_format = template_format.heading3 if template_format else None
    if text_format is not None:
        paragraph = document.add_paragraph()
        _apply_text_format(paragraph, text_format)
    else:
        paragraph = document.add_heading("", level=3)
        paragraph.paragraph_format.first_line_indent = Pt(0)
    problem_text = _format_problem_text(problem_text)
    run = paragraph.add_run(heading_text)
    if text_format is not None:
        _apply_run_format(run, text_format)
    else:
        _set_run_font(run, style.body_font, style.body_size_pt)
    if problem_text:
        problem_run = paragraph.add_run(problem_text)
        if text_format is not None:
            _apply_run_format(problem_run, text_format)
        else:
            _set_run_font(problem_run, style.body_font, style.body_size_pt)
        return

    warning_run = paragraph.add_run("【缺少问题内容】")
    _set_run_font(warning_run, style.body_font, style.body_size_pt, True)
    warning_run.font.color.rgb = RGBColor(255, 0, 0)


def _format_problem_text(problem_text: str) -> str:
    """把点位后的具体问题格式化为连续编号文本。"""

    text = _normalize_sentence_end(" ".join(str(problem_text).splitlines()).strip())
    if not text or text in PLAIN_PROBLEM_TEXTS:
        return text

    parts = _split_problem_parts(text)
    if not parts:
        return ""

    formatted = []
    for index, part in enumerate(parts, start=1):
        ending = "。" if index == len(parts) else "；"
        formatted.append(f"（{index}）{part}{ending}")
    return "".join(formatted)


def _strip_problem_marker(text: str) -> str:
    """去掉原文本中已有的序号和结尾标点，避免重复编号。"""

    cleaned = text.strip()
    cleaned = re.sub(r"^[（(]\d+[）)]\s*", "", cleaned)
    cleaned = re.sub(r"^\d{1,2}[、.．]\s*", "", cleaned)
    return _normalize_sentence_end(cleaned).strip(" ；;。.")


def _split_problem_parts(text: str) -> List[str]:
    """按标点和原始数字序号拆分问题项。"""

    normalized = re.sub(r"(?<!^)(?<![（(])(?<!\d)(\d{1,2})[、.．]", r"；\1、", text)
    parts = []
    for part in re.split(r"[;；。]+", normalized):
        cleaned = _strip_problem_marker(part)
        if cleaned:
            parts.append(cleaned)
    return parts


def _normalize_sentence_end(text: str) -> str:
    """压缩重复句号，避免输出结尾出现两个句号。"""

    return re.sub(r"[。.]{2,}$", "。", text.strip())


def _add_images(document: Document, image_paths: Iterable[str], style: ReportStyle) -> None:
    """在同一个段落中连续插入图片，不使用表格，也不插入空格。"""

    paths = list(image_paths)[: style.pictures_per_row]
    if not paths:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for image_path in paths:
        run = paragraph.add_run()
        try:
            run.add_picture(image_path, width=Cm(style.image_width_cm), height=Cm(style.image_height_cm))
        except Exception:
            run.add_text("[图片插入失败]")


def _add_statistics_table(
    document: Document,
    rows: List[Dict[str, Any]],
    style: ReportStyle,
    template_format: TemplateFormat | None = None,
) -> None:
    columns = ["街道名称", *indicators_for("transfer_station"), "问题总数"]
    table = document.add_table(rows=1, cols=len(columns))
    if template_format and template_format.table_style_name:
        try:
            table.style = template_format.table_style_name
        except KeyError:
            pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, column in enumerate(columns):
        _set_cell_text(
            table.rows[0].cells[i],
            column,
            style.body_font,
            style.table_size_pt,
            True,
            text_format=template_format.table_header if template_format else None,
        )
    for row_data in rows:
        cells = table.add_row().cells
        for i, column in enumerate(columns):
            value = row_data.get(column, 0)
            if isinstance(value, int):
                value = style.table_zero_text if value == 0 else str(value)
            _set_cell_text(
                cells[i],
                value,
                style.body_font,
                style.table_size_pt,
                text_format=template_format.table_body if template_format else None,
            )
    _set_table_borders(table)


def _template_title(document: Document) -> str:
    """读取模板正文中的首个非空段落作为标题文本。"""

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            return text
    return ""


def _new_document(template_path: Path | None, style: ReportStyle) -> tuple[Document, str, TemplateFormat | None, TemplateSections | None]:
    """创建报告文档；有模板时继承模板格式，无模板时使用默认页面。"""

    if template_path:
        document = Document(str(template_path))
        title = _template_title(document)
        template_format = _capture_template_format(document)
        template_sections = _capture_template_sections(document)
        _clear_body_keep_section(document, template_sections.first_section_pr)
        return document, title, template_format, template_sections

    document = Document()
    _set_a4_section(document.sections[0], WD_ORIENT.PORTRAIT, style)
    return document, "", None, None


def _cn_index(index: int) -> str:
    """把 1-based 序号转为简单中文序号。"""

    if 1 <= index <= len(CHINESE_NUMBERS):
        return CHINESE_NUMBERS[index - 1]
    if index < 20:
        return "十" + CHINESE_NUMBERS[index - 11]
    return str(index)


def _group_records(records: List[Dict[str, Any]]) -> "OrderedDict[str, List[Dict[str, Any]]]":
    grouped: "OrderedDict[str, List[Dict[str, Any]]]" = OrderedDict()
    for record in records:
        group = record.get("report_group") or record.get("street") or "未识别街道"
        grouped.setdefault(group, []).append(record)
    return grouped


def _merge_point_records(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """同一检查点在案例正文中只输出一次，多个问题合并后统一编号。"""

    merged: "OrderedDict[str, Dict[str, Any]]" = OrderedDict()
    for record in records:
        point_name = str(record.get("report_point") or record.get("location") or "未识别点位").strip()
        key = point_name or "未识别点位"
        if key not in merged:
            merged[key] = {**record, "report_point": point_name, "specific_problem": "", "images": []}

        target = merged[key]
        target["has_problem"] = bool(target.get("has_problem")) or bool(record.get("has_problem"))
        _append_problem_text(target, _problem_text(record))
        target.setdefault("images", []).extend(record.get("images", []))

    return list(merged.values())


def _append_problem_text(record: Dict[str, Any], problem_text: str) -> None:
    text = _normalize_sentence_end(" ".join(str(problem_text).splitlines()).strip())
    if not text:
        return

    current = str(record.get("specific_problem") or "").strip()
    if not current:
        record["specific_problem"] = "；".join(_split_problem_parts(text)) if text not in PLAIN_PROBLEM_TEXTS else text
        return
    if current in PLAIN_PROBLEM_TEXTS and text not in PLAIN_PROBLEM_TEXTS:
        record["specific_problem"] = text
        return
    if text in PLAIN_PROBLEM_TEXTS:
        return

    existing_parts = set(_split_problem_parts(current))
    next_parts = _split_problem_parts(text)
    for part in next_parts:
        if part not in existing_parts:
            current = f"{current}；{part}" if current else part
            existing_parts.add(part)
    record["specific_problem"] = current


def _problem_text(record: Dict[str, Any]) -> str:
    if record.get("has_problem"):
        text = str(record.get("secondary_indicator") or "").strip()
        if text:
            return text
        indicators = record.get("indicators") or []
        if indicators:
            return "；".join(str(indicator).strip() for indicator in indicators if str(indicator).strip())

    text = str(record.get("specific_problem") or record.get("problem") or "").strip()
    if text:
        return text
    return "" if record.get("has_problem") else "无问题"


def generate_report(
    records: List[Dict[str, Any]],
    stats: Dict[str, Any],
    output_path: Path,
    title: str,
    start_date: str,
    end_date: str,
    style: ReportStyle = REPORT_STYLE,
    template_path: Path | None = None,
) -> Path:
    """根据记录和统计结果生成 Word 月报。"""

    if template_path is not None and output_path.resolve() == template_path.resolve():
        raise ValueError("输出文件不能覆盖模板文件，请更换输出文件名或输出目录。")

    document, template_title, template_format, template_sections = _new_document(template_path, style)

    _add_title(document, template_title or title, style, template_format)
    _add_heading(document, "一、总体情况", level=1, template_format=template_format)
    _add_body_paragraph(document, stats["summary_text"], style, template_format=template_format)
    _add_heading(document, "二、各街道案例", level=1, template_format=template_format)

    for group_index, (group_name, group_records) in enumerate(_group_records(records).items(), start=1):
        _add_heading(document, f"（{_cn_index(group_index)}）{group_name}", level=2, template_format=template_format)
        for point_index, record in enumerate(_merge_point_records(group_records), start=1):
            point_name = record.get("report_point") or record.get("location") or "未识别点位"
            _add_point_problem_paragraph(
                document,
                f"{point_index}.{point_name}可回收物交投点：",
                _problem_text(record),
                style,
                template_format=template_format,
            )
            _add_images(document, record.get("images", []), style)

    _add_landscape_section(
        document,
        style,
        keep_template_margins=template_format is not None,
        template_sections=template_sections,
    )
    _add_heading(document, f"附件：{start_date}至{end_date}可回收物交投点各项指标问题数", level=1, template_format=template_format)
    _add_statistics_table(document, stats["table_rows"], style, template_format=template_format)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
