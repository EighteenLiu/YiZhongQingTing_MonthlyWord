"""厨余垃圾就地处理 Jinja 模板报告生成。"""

from __future__ import annotations

from collections import OrderedDict, defaultdict
from datetime import date, datetime
import re
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Callable, Dict, Iterable, List

from docx.shared import Cm


KITCHEN_ITEMS = [
    "公示牌",
    "站点周边环境",
    "厨余垃圾处理记录台账",
    "厨余尾料废弃物处置台账",
    "安全检查记录台账",
    "消防演练资料",
    "安全守则、操作流程、处理工艺流程、处理设备组成",
    "灭蝇灯",
    "灭火器",
]
SPECIAL_STATUS_KEYWORDS = [
    "已拆除",
    "已停运",
    "停运",
    "停用",
    "锁门",
    "负责人",
    "失联",
    "无法入场",
    "无法进入",
    "无处理设备",
    "暂存点",
]
NO_PROBLEM_TEXTS = {"", "1", "无问题", "无问题。", "正常", "合格", "良好，未发现问题"}
DEFAULT_DEPARTMENT = "北京市西城区城市管理委员会"
DEFAULT_SUMMARY_TEXT = "各点位存在问题，请相关单位尽快完成整改，并将整改情况报送至区城管委固废科。"
CN_NUMBERS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


ImageFactory = Callable[[str], Any]


def _text(value: Any) -> str:
    return str(value or "").strip()


def _compact(value: Any) -> str:
    return re.sub(r"\s+", "", _text(value)).strip("。；;，,")


def _cn_index(index: int) -> str:
    if 1 <= index <= len(CN_NUMBERS):
        return CN_NUMBERS[index - 1]
    if index < 20:
        return "十" + CN_NUMBERS[index - 11]
    return str(index)


def _to_attr_object(value: Any) -> Any:
    if isinstance(value, dict):
        return SimpleNamespace(**{key: _to_attr_object(item) for key, item in value.items()})
    if isinstance(value, list):
        return [_to_attr_object(item) for item in value]
    return value


def _photo_rows(image_paths: Iterable[str], image_factory: ImageFactory | None = None) -> List[Dict[str, Any]]:
    images = [image_factory(path) if image_factory else path for path in image_paths]
    rows = []
    for index in range(0, len(images), 2):
        rows.append(
            {
                "left": images[index],
                "right": images[index + 1] if index + 1 < len(images) else "",
            }
        )
    return rows


def _ensure_status_photo_template(template_path: Path) -> None:
    """给旧版厨余模板补充特殊状态行的照片渲染块。"""

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    document = Document(template_path)
    if any("station.status_photo_rows" in paragraph.text for paragraph in document.paragraphs):
        return

    else_match = next(
        ((index, paragraph) for index, paragraph in enumerate(document.paragraphs) if "{%p else %}" in paragraph.text),
        None,
    )
    if else_match is None:
        return

    else_index, else_paragraph = else_match
    base_style = document.paragraphs[else_index - 1].style if else_index > 0 else else_paragraph.style
    insertions = [
        ("{%p if station.status_photo_rows %}", None),
        ("{%p for row in station.status_photo_rows %}", None),
        ("{{ row.left }}{{ row.right }}", WD_ALIGN_PARAGRAPH.CENTER),
        ("{%p endfor %}", None),
        ("{%p endif %}", None),
    ]
    for text, align in insertions:
        paragraph = else_paragraph.insert_paragraph_before(text)
        paragraph.style = base_style
        if align is not None:
            paragraph.alignment = align

    document.save(template_path)


def _raw_value(record: Dict[str, Any], key: str) -> str:
    raw = record.get("raw") or {}
    return _text(raw.get(key))


def _station_name(record: Dict[str, Any]) -> str:
    return _text(record.get("report_point")) or _text(record.get("location")) or _raw_value(record, "4级点位") or "未识别站点"


def _station_visit_date(record: Dict[str, Any]) -> str:
    raw_date = _raw_value(record, "案件上报时间") or _text(record.get("check_date"))
    return raw_date.split()[0] if raw_date else ""


def _station_visit_datetime(record: Dict[str, Any]) -> datetime | None:
    raw_date = _raw_value(record, "案件上报时间") or _text(record.get("check_date"))
    if not raw_date:
        return None

    candidates = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
        "%Y年%m月%d日 %H:%M:%S",
        "%Y年%m月%d日 %H:%M",
        "%Y年%m月%d日",
    ]
    for fmt in candidates:
        try:
            return datetime.strptime(raw_date.strip(), fmt)
        except ValueError:
            continue
    return None


def _station_visit_day(record: Dict[str, Any]) -> date | None:
    visit_dt = _station_visit_datetime(record)
    return visit_dt.date() if visit_dt else None


def _latest_records_by_station(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """同一站点多天检查时，仅保留最新检查日期的记录。"""

    latest_by_station: Dict[str, date | None] = {}
    for record in records:
        station_name = _station_name(record)
        visit_dt = _station_visit_day(record)
        current = latest_by_station.get(station_name)
        if current is None or (visit_dt is not None and visit_dt > current):
            latest_by_station[station_name] = visit_dt

    filtered: List[Dict[str, Any]] = []
    for record in records:
        station_name = _station_name(record)
        visit_dt = _station_visit_day(record)
        latest = latest_by_station.get(station_name)
        if latest is None or visit_dt == latest:
            filtered.append(record)
    return filtered


def _item_name(record: Dict[str, Any]) -> str:
    secondary = _text(record.get("secondary_indicator")) or _raw_value(record, "2级指标")
    if secondary in KITCHEN_ITEMS:
        return secondary
    source = secondary
    for item in KITCHEN_ITEMS:
        if item in source:
            return item
    return ""


def _issue_item_names(record: Dict[str, Any], issue_text: str) -> List[str]:
    source = f"{_item_name(record)} {issue_text}"
    aliases = {
        "厨余垃圾处理记录台账": ["厨余垃圾处理记录", "处理记录台账"],
        "厨余尾料废弃物处置台账": ["厨余尾料", "尾料废弃物", "尾料处置"],
        "安全检查记录台账": ["安全检查记录", "安全检查台账"],
        "消防演练资料": ["消防演练资料", "消防培训演练", "消防演练"],
        "安全守则、操作流程、处理工艺流程、处理设备组成": ["安全守则", "操作流程", "处理工艺流程", "处理设备组成"],
        "灭蝇灯": ["灭蝇灯"],
        "灭火器": ["灭火器"],
        "公示牌": ["公示牌"],
        "站点周边环境": ["周边环境", "站点周边"],
    }
    matched = [item for item, keywords in aliases.items() if item in source or any(keyword in source for keyword in keywords)]
    primary = _item_name(record)
    if primary and primary not in matched:
        matched.insert(0, primary)
    return matched


def _status_text(record: Dict[str, Any]) -> str:
    secondary = _text(record.get("secondary_indicator")) or _raw_value(record, "2级指标")
    raw_problem = _raw_value(record, "具体问题")
    source = f"{secondary} {raw_problem}"
    if any(keyword in source for keyword in SPECIAL_STATUS_KEYWORDS):
        return secondary if secondary not in KITCHEN_ITEMS else raw_problem
    return ""


def _issue_text(record: Dict[str, Any]) -> str:
    raw_problem = _raw_value(record, "具体问题")
    if _compact(raw_problem) in {_compact(text) for text in NO_PROBLEM_TEXTS}:
        return ""
    return raw_problem


def _append_unique(parts: List[str], value: str) -> None:
    text = value.strip()
    if text and text not in parts:
        parts.append(text)


def _issue_date_from_report_month(report_month: str) -> str:
    text = _text(report_month)
    match = re.match(r"(\d{4})年(\d{1,2})月", text)
    if match:
        return f"{match.group(1)}年{int(match.group(2))}月22日"
    try:
        parsed = datetime.strptime(text, "%Y-%m").date()
        return f"{parsed.year}年{parsed.month}月22日"
    except ValueError:
        today = date.today()
        return f"{today.year}年{today.month}月{today.day}日"


def _display_name(name: str, visit_dates: set[str]) -> str:
    return f"{name}厨余垃圾就地处理站点"


def _problem_text(stations: List[Dict[str, Any]]) -> str:
    status_groups: Dict[str, List[str]] = defaultdict(list)
    missing_groups: Dict[str, List[str]] = defaultdict(list)

    for station in stations:
        name = station["name"]
        status = _text(station.get("status_summary"))
        if status:
            status_groups[status].append(name)
            continue
        for item in station.get("items", []):
            if item.get("has_issue"):
                missing_groups[item["name"]].append(name)

    parts = []
    for status, names in status_groups.items():
        parts.append(f"{'、'.join(names)}等{len(names)}个站点{status}")
    if missing_groups:
        detail = "；".join(f"{'、'.join(names)}缺少{item}" for item, names in missing_groups.items())
        parts.append(f"运营中的站点存在资料或设施缺失问题，{detail}")
    return "。".join(parts) + ("。" if parts else "未发现明显问题。")


def build_kitchen_waste_context(
    records: List[Dict[str, Any]],
    report_month: str,
    title: str = "西城区厨余垃圾就地处理检查报告",
    image_factory: ImageFactory | None = None,
) -> Dict[str, Any]:
    """把清洗后的厨余记录转换为模板所需上下文。"""

    records = _latest_records_by_station(records)
    grouped: "OrderedDict[str, Dict[str, Any]]" = OrderedDict()
    for record in records:
        station_name = _station_name(record)
        if station_name not in grouped:
            grouped[station_name] = {
                "name": station_name,
                "visit_dates": set(),
                "statuses": [],
                "status_images": [],
                "items": {
                    item_name: {"index_cn": _cn_index(index), "name": item_name, "issues": [], "images": [], "seen": False}
                    for index, item_name in enumerate(KITCHEN_ITEMS, start=1)
                },
            }

        station = grouped[station_name]
        visit_date = _station_visit_date(record)
        if visit_date:
            station["visit_dates"].add(visit_date)

        status = _status_text(record)
        if status:
            _append_unique(station["statuses"], status)
            station["status_images"].extend(record.get("images", []))
            continue

        item_name = _item_name(record)
        if not item_name:
            continue
        item = station["items"][item_name]
        item["seen"] = True
        item["images"].extend(record.get("images", []))

        issue = _issue_text(record)
        if issue:
            for issue_item_name in _issue_item_names(record, issue):
                issue_item = station["items"][issue_item_name]
                issue_item["seen"] = True
                _append_unique(issue_item["issues"], "无")

    stations = []
    for station in grouped.values():
        status_summary = "，".join(station["statuses"])
        if status_summary:
            stations.append(
                {
                    "name": station["name"],
                    "display_name": _display_name(station["name"], station["visit_dates"]),
                    "status_summary": status_summary,
                    "status_photo_rows": _photo_rows(station["status_images"], image_factory=image_factory),
                    "items": [],
                }
            )
            continue

        items = []
        for item in station["items"].values():
            has_images = bool(item["images"])
            result_text = "" if has_images else "无"
            items.append(
                {
                    "index_cn": item["index_cn"],
                    "name": item["name"],
                    "result_text": result_text,
                    "has_issue": bool(item["issues"]),
                    "photo_rows": _photo_rows(item["images"], image_factory=image_factory),
                }
            )
        stations.append(
                {
                    "name": station["name"],
                    "display_name": _display_name(station["name"], station["visit_dates"]),
                    "status_summary": "",
                    "status_photo_rows": [],
                    "items": items,
                }
            )

    return {
        "title": title or "西城区厨余垃圾就地处理检查报告",
        "report_month_text": report_month,
        "site_names_text": "、".join(station["name"] for station in stations),
        "department": DEFAULT_DEPARTMENT,
        "stations": stations,
        "problem_text": _problem_text(stations),
        "summary_text": DEFAULT_SUMMARY_TEXT,
        "issue_date_text": _issue_date_from_report_month(report_month),
    }


def summarize_kitchen_waste_records(records: List[Dict[str, Any]], report_month: str = "") -> Dict[str, Any]:
    """厨余报告专用统计口径。"""

    context = build_kitchen_waste_context(records, report_month)
    problem_count = 0
    for station in context["stations"]:
        if station["status_summary"]:
            problem_count += 1
        problem_count += sum(1 for item in station["items"] if item.get("has_issue"))

    return {
        "street_count": len(context["stations"]),
        "site_count": len(context["stations"]),
        "record_count": len(records),
        "problem_record_count": problem_count,
        "problem_count": problem_count,
        "summary_text": (
            f"{context['report_month_text']}区城管委固废科对"
            f"{context['site_names_text']}厨余垃圾就地处理站点进行现场检查。"
        ),
        "table_rows": [],
    }


def generate_kitchen_waste_report(
    records: List[Dict[str, Any]],
    output_path: Path,
    title: str,
    report_month: str,
    template_path: Path,
) -> Path:
    """使用 docxtpl 渲染厨余垃圾就地处理检查报告。"""

    if output_path.resolve() == template_path.resolve():
        raise ValueError("输出文件不能覆盖模板文件，请更换输出文件名或输出目录。")

    try:
        from docxtpl import DocxTemplate, InlineImage
    except ImportError as exc:
        raise RuntimeError("缺少 docxtpl/jinja2 依赖，请运行 run_app.bat 重新安装 requirements.txt。") from exc

    _ensure_status_photo_template(template_path)
    template = DocxTemplate(str(template_path))

    def image_factory(path: str) -> Any:
        return InlineImage(template, path, width=Cm(7.21), height=Cm(4.06))

    context = build_kitchen_waste_context(
        records=records,
        report_month=report_month,
        title=title,
        image_factory=image_factory,
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    template.render({key: _to_attr_object(value) for key, value in context.items()})
    template.save(output_path)
    return output_path
